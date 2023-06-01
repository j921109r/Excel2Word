'''
Author: Garrick
Date: 2022-09-05 09:21:38
LastEditors: Garrick
LastEditTime: 2022-09-13 09:48:01
Descripttion: This is for ...
'''
import xlrd
import docx
from docx import Document
from docx.shared import Pt,Cm
from docxcompose.composer import Composer
import os
import re


def open_excel(file= 'file.xls'):  
    try:  
        data = xlrd.open_workbook(file)  
        return data  
    except Exception as e:  
        print(str(e))  
        
def excel_table_byname(file= 'E:\\個人檔案\\6-desktop\\xxxx.xls',colnameindex=0,by_name=u'電度'):  
    data = open_excel(file)  
    table = data.sheet_by_name(by_name)  
    nrows = table.nrows #行數   
    colnames =  table.row_values(colnameindex) #某一行資料   
    list =[]  
    for rownum in range(1,nrows):  
         row = table.row_values(rownum)  
         if row:  
             app = {}  
             for i in range(len(colnames)):  
                app[colnames[i]] = row[i]  
             list.append(app)  
    return list  
  
def main(name): 
    data = excel_table_byname(name+".xls",1,u'Sheet0')
    num = int(re.findall(r'-?\d+\.?\d*', name)[0])
    # print (data[0]['個案姓名'])
    path = f'C:\\Users\\j9211\\Desktop\\Excel2Word\\{name}'
    if not os.path.isdir(path):
        os.mkdir(path)
    doc = docx.Document('1.docx')
    times = int((num/14)+1)
    res = (num)%14
    for j in range(times):
        table = doc.tables[0]
        if (j==times-1):
            for i in range(0, len(table.rows)):
                base = (i*2)+(14*j) 
                if (((i*2)+0)>res):
                    table.cell(i,0).text = ""
                    table.cell(i,1).text = ""
                elif(((i*2)+1)>res):
                    table.cell(i,1).text = ""
                else:
                    table.cell(i,0).paragraphs[0].text = ""+data[base+1]["個案姓名"]
                    table.cell(i,0).paragraphs[1].text = ""+data[base+1]["戶籍村里"]+data[base+1]["戶籍地址"]
                    paragraphs = table.cell(i,0).paragraphs[0]
                    paragraphs.runs[0].bold = True
                    paragraphs.runs[0].font.size = Pt(20)
                    paragraphs = table.cell(i,0).paragraphs[1]
                    paragraphs.runs[0].bold = True
                    if (res%2 ==1)&(((i*2)+1)>=res):
                        pass
                    else:
                        table.cell(i,1).paragraphs[1].text = ""+data[base+2]["戶籍村里"]+data[base+2]["戶籍地址"]  
                        table.cell(i,1).paragraphs[0].text = ""+data[base+2]["個案姓名"]
                        paragraphs = table.cell(i,1).paragraphs[0]
                        paragraphs.runs[0].bold = True
                        paragraphs.runs[0].font.size = Pt(20)
                        paragraphs = table.cell(i,1).paragraphs[1]
                        paragraphs.runs[0].bold = True
        else:
            for i in range(0, len(table.rows)):
                base = (i*2)+(14*j) 
                table.cell(i,0).paragraphs[0].text = ""+data[base+1]["個案姓名"]
                table.cell(i,0).paragraphs[1].text = ""+data[base+1]["戶籍村里"]+data[base+1]["戶籍地址"]
                table.cell(i,1).paragraphs[0].text = ""+data[base+2]["個案姓名"]
                table.cell(i,1).paragraphs[1].text = ""+data[base+2]["戶籍村里"]+data[base+2]["戶籍地址"]  
                paragraphs = table.cell(i,0).paragraphs[0]
                paragraphs.runs[0].bold = True
                paragraphs.runs[0].font.size = Pt(20)
                paragraphs = table.cell(i,0).paragraphs[1]
                paragraphs.runs[0].bold = True
                paragraphs = table.cell(i,1).paragraphs[0]
                paragraphs.runs[0].bold = True
                paragraphs.runs[0].font.size = Pt(20)
                paragraphs = table.cell(i,1).paragraphs[1]
                paragraphs.runs[0].bold = True

        doc.save(f'{path}\{name}-{j+1}.docx')
    
        
if __name__=="__main__":  
    # list1 = os.listdir()
    # for i in list1:
    #     if (".xls" in i ):
    #         pass
    #     else:
    #         list1.remove(i)
    # for i in list1:
    #     if (".xls" in i ):
    #         pass
    #     else:
    #         list1.remove(i)
    # for i in list1:
    #     if (".xls" in i ):
    #         pass
    #     else:
    #         list1.remove(i)
    # # print(list1)
    # for item in list1:
    #     name = item.split(".",1)
    #     main(name[0])  
    #     print(name[0]+" OK")
    # main('英仁里30人')
    
    New_doc = Document()
    section = New_doc.sections[0]
    section.top_margin=Cm(1.27)
    
    com_doc = Composer(New_doc)
    
    doc2 = Document('C:/Users/j9211/Desktop/backup/Excel2Word/Page2.docx')
    # doc2.add_page_break()
    
    data = excel_table_byname("COVID-19名冊-已打一劑.xls",1,u'Sheet0')
    for data_detail in data:
        # New template
        doc = Document('C:/Users/j9211/Desktop/backup/Excel2Word/Page1.docx')
        doc.add_page_break()
        # bypass title
        if (data_detail['證號']=='證號'):
            continue
        # search replacement
        for para in doc.paragraphs:
            for i in range(len(para.runs)):
                # print(para.runs[i].text)
                # replace name
                if('袁子英'in para.runs[i].text):
                    para.runs[i].text = para.runs[i].text.replace('袁子英', data_detail['個案姓名'])
                # replace address
                if('吉仁里3鄰南榮路'in para.runs[i].text):
                    para.runs[i].text = para.runs[i].text.replace('吉仁里3鄰南榮路', data_detail['戶籍地址'])    

        # add to com_doc
        com_doc.append(doc)
        com_doc.append(doc2)
    
        # print(para.text)


    # doc = Document('C:/Users/j9211/Desktop/backup/Excel2Word/example.docx')
    # doc.add_page_break()
    # # add to com_doc
    # com_doc.append(doc)
    # com_doc.append(doc2)
    
    
    com_doc.save(f'C:/Users/j9211/Desktop/backup/Excel2Word/example3.docx')
    pass