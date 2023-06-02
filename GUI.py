
from PyQt5 import QtWidgets, QtCore, QtGui
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import *
# from PyQt5.QtGui import QAction

import xlrd
import docx
from docx import Document
from docx.shared import Pt,Cm
from docxcompose.composer import Composer
from docx2pdf import convert
import win32com.client

from Main_ui import Ui_MainWindow
import traceback,sys
from python import *

class MainUi(QtWidgets.QMainWindow, Ui_MainWindow): 
    def __init__(self):
        super().__init__()
        QtWidgets.QMainWindow.__init__(self)
        Ui_MainWindow.__init__(self)
        self.setupUi(self)
        self.init_connect()
        self.init_value()
        self.center()
        
    def init_value(self):  
        self.MCU_REQ_path=''  
        
    def init_connect(self):
        self.Button_Excel.clicked.connect(self.click_Button_Excel)    
        self.Button_GO.clicked.connect(self.click_Button_Go)    
        
    def center(self):
        # 獲得主視窗所在的框架
        qr = self.frameGeometry()
        # 獲取顯示器的解析度，然後得到螢幕中間點的位置
        cp = QtGui.QGuiApplication.primaryScreen().availableGeometry().center()
        # 然後把主視窗框架的中心點放置到螢幕的中心位置
        qr.moveCenter(cp)
        # 然後通過 move 函式把主視窗的左上角移動到其框架的左上角
        self.move(qr.topLeft())        
        
        
    def click_Button_Excel(self):
        dialog = QtWidgets.QFileDialog()
        Excel_path=dialog.getOpenFileName(self,"choose File",None,'Xls (*.xls)')
        self.Excel_path=Excel_path[0]
        self.Label_Excel.setText(self.Excel_path)
        self.textBrowser.setText("Idel .....")
        pass
    
    def click_Button_Go(self):
        if self.Excel_path=='':
            pass
        
        New_doc = Document()
            
            
        section = New_doc.sections[0]
        section.top_margin=Cm(1.27)

        com_doc = Composer(New_doc)

        doc2 = Document('./Page2.docx')
        # doc2.add_page_break()

        data = excel_table_byname(self.Excel_path,1,u'Sheet0')
        for data_detail in data:
            # New template
            doc = Document('./Page1.docx')
            doc.add_page_break()
            # bypass title
            if (data_detail['證號']=='證號'):
                continue
            # search replacement
            for para in doc.paragraphs:
                for i in range(len(para.runs)):
                    # print(para.runs[i].text)
                    # replace name
                    if('袁子英'in para.runs[i].text):
                        para.runs[i].text = para.runs[i].text.replace('袁子英', data_detail['個案姓名'])
                    # replace address
                    
                    if('吉仁里3鄰南榮路'in para.runs[i].text):
                        if ('戶籍村里'in data[0].keys()):                        
                            para.runs[i].text = para.runs[i].text.replace('吉仁里3鄰南榮路', data_detail['戶籍村里']+data_detail['戶籍地址'])    
                        else:
                            para.runs[i].text = para.runs[i].text.replace('吉仁里3鄰南榮路', data_detail['戶籍地址'])    

            # add to com_doc
            com_doc.append(doc)
            com_doc.append(doc2)
            
            output_name = self.lineEdit_name.text()
            if output_name =='':
                output_name=f'{os.path.splitext(os.path.basename(self.Excel_path))[0]}'
            
            com_doc.save(f'./{output_name}.docx')
        
        if (self.checkBox_PDF.isChecked()):
            file = open(f'./{output_name}.pdf', "w")
            file.close()
            convert(f'./{output_name}.docx',f'./{output_name}.pdf')
        # word = win32com.client.Dispatch('Word.Application')
        # doc = word.Documents.Open(f'./{output_name}.docx')
        # wdFormatPDF = 17
        # doc.SaveAs(f'./{output_name}.pdf', FileFormat=wdFormatPDF)
        # doc.Close()
        # word.Quit()
        self.textBrowser.setText("OK~")
        pass
    
    
def run_app():
    app = QtWidgets.QApplication(sys.argv)
    controller = MainUi()
    controller.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    try:
        run_app()
    except Exception:
        print (traceback.format_exc())